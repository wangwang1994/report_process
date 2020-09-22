from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt,RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from tkinter import *
from docx.oxml.ns import qn
import re
import pprint
import openpyxl
import os
with open(os.getcwd()+'/chuli.txt',encoding="gbk") as f:
    text=f.read()
dict={}
#print(text)
text=text.strip()
text=text.replace('\n','')
text=text.replace('\r','')
text=text.replace(' ','')
print(text)


print('---------车辆型号------------')
cheliangxinghao_regex=re.compile(r'1(车辆型号：)(.*)2商标')
cheliangxinghao=cheliangxinghao_regex.search(text)
# print(cheliangxinghao.group(1))           #得到的group是str类型的文本了
# print(cheliangxinghao.group(2))
dict[cheliangxinghao.group(1)]=cheliangxinghao.group(2)
# print(dict
# print(type(cheliangxinghao.group(1)))
print('---------商标：------------')
brand_regex=re.compile(r'2(商标：)(.*)3汽车分类')
brand=brand_regex.search(text)
# print(brand.group(1))           #得到的group是str类型的文本了
# print(brand.group(2))
dict['商标']=brand.group(2)
# print(dict)
print('------------汽车分类----------------')
qichefenlei_regex=re.compile(r'\d(汽车分类：)(\w{2})\d')
qichefenlei=qichefenlei_regex.search(text)
# print(qichefenlei.group(0))
dict[qichefenlei.group(1)]=qichefenlei.group(2)
# print(dict)
print('-----------车辆制造商名称----------------')
oem_regex=re.compile(r'\d(车辆制造商名称：)(.*)7生')
oem=oem_regex.search(text)
# print(oem.group())
dict[oem.group(1)]=oem.group(2)
# print(dict)
print('-----------生产厂地址：----------------')
address_regex=re.compile(r'\d(生产厂地址：)(.*)\d发动机编号')
address=address_regex.search(text)
# print(address.group())
dict[address.group(1)]=address.group(2)
# print(dict)
print('-----------发动机型号/生产企业：----------------')
type_eng_regex=re.compile(r'\d(发动机型号/生产企业：)(.*?)1\d催化转化器')
type_eng=type_eng_regex.search(text)
# print(type_eng.group())
dict[type_eng.group(1)]=type_eng.group(2)
# print(dict)
print('-----------发动机型号/生产企业：单独的----------------')
type_eng1_regex=re.compile(r'\d(发动机型号/生产企业：)(.*?)/(.*?)1\d')
type_eng1=type_eng1_regex.search(text)
# print(type_eng.group())
dict[type_eng1.group(1)+'1']=type_eng1.group(2)
dict[type_eng1.group(1)+'2']=type_eng1.group(3)
print('-----------vin：----------------')
vin_eng_regex=re.compile(r'kg(\w*)（VIN码）')
vin=vin_eng_regex.search(text)
# print(type_eng.group())
dict['vin']=type_eng.group(1)
print('-----------发动机编号：发动机编号：----------------')
serial_regex=re.compile(r'\d(发动机编号：)(\w*)\d')
serial=serial_regex.search(text)
# print(serial.group())
dict[serial.group(1)]=serial.group(2)
# print(dict)
print('-----------基准质量：----------------')
mass_regex=re.compile(r'\d(基准质量：)(\w*)kg')
mass=mass_regex.search(text)
# print(mass.group())
dict[mass.group(1)]=mass.group(2)
# print(dict)
print('-----------催化转化器型号/生产企业：----------------')
catalyst_regex=re.compile(r'催化转化器型号/生产企业：(.*?)(涂层)')
catalyst=catalyst_regex.search(text)
# print(catalyst.group())
dict['催化转化器']=catalyst.group(1)
# print(dict)
print('-----------催化转化器型号/生产企业：详细拆分----------------')
twcparts_regex=re.compile(r'\d\d(催化转化器型号/生产企业：)(.*?)/前:')
twcparts=twcparts_regex.search(text)
dict['催化转化器型号']=twcparts.group(2)
# dict['催化转化器生产企业']=twcparts.group(3)
# print(catalyst.group())
# dict[catalyst.group(1)]=catalyst.group(2)
# print(dict)
print('-----------twc涂层/载体/封装生产企业：----------------')
catalyst_layer_regex=re.compile(r'(涂层/载体/封装生产企业：)(.*?)\d\d')
catalyst_layer=catalyst_layer_regex.search(text)
# print(catalyst_layer.group())
dict['催化转化器涂层']=catalyst_layer.group(2)
# pprint.pprint(dict)
print('-----------颗粒捕集器型号/生产企业：----------------')
dpf_regex=re.compile(r'\d\d(颗粒捕集器型号/生产企业：)(.*?)(涂层)')
dpf=dpf_regex.search(text)
# print(dpf.group())
dict['颗粒捕集器']=dpf.group(2)
# pprint.pprint(dict)
print('-----------dpf涂层/载体/封装生产企业：----------------')
dpf_layer_regex=re.compile(r'(颗粒捕集器型号/生产企业：)(.*?)(涂层/载体/封装生产企业：)(.*?)\d\d')
dpf_layer=dpf_layer_regex.search(text)
# print(dpf_layer.group())
dict['颗粒捕集器涂层']=dpf_layer.group(4)
# pprint.pprint(dict)
print('----------- 炭罐型号/生产企业：----------------')
carbon_regex=re.compile(r'\d\d(炭罐型号/生产企业：)(.*)(\d\d)(氧传感器)')
carbon=carbon_regex.search(text)
# print(carbon.group())
dict[carbon.group(1)]=carbon.group(2)
# pprint.pprint(dict)
print('----------- 氧传感器型号/生产企业：----------------')
ox_regex=re.compile(r'\d\d(氧传感器型号/生产企业：)(.*)(\d\d)(曲轴箱)')
ox=ox_regex.search(text)
# print(carbon.group())
dict[ox.group(1)]=ox.group(2)
print('----------- 曲轴箱排放控制装置型号/生产企业：----------------')
crank_regex=re.compile(r'\d\d(曲轴箱排放控制装置型号/生产企业：)(.*)(\d\d)(EGR)')
crank=crank_regex.search(text)
# print(crank.group())
dict[crank.group(1)]=crank.group(2)
# pprint.pprint(dict)

print('----------- EGR型号/生产企业：----------------')
egr_regex=re.compile(r'\d\d(EGR型号/生产企业：)(.*)(\d\d)(OBD系统)')
egr=egr_regex.search(text)
# print(egr.group())
dict[egr.group(1)]=egr.group(2)
# pprint.pprint(dict)

print('----------- OBD系统供应商：----------------')
obd_regex=re.compile(r'\d\d(OBD系统供应商：)(.*)(\d\d)(ECU)')
obd=obd_regex.search(text)
# print(obd.group())
dict[obd.group(1)]=obd.group(2)
# pprint.pprint(dict)

print('----------- ECU型号/生产企业：----------------')
ecu_regex=re.compile(r'\d\d(ECU型号/生产企业：)(.*)(\d\d)(变速器型式)')
ecu=ecu_regex.search(text)
# print(ecu.group())
dict[ecu.group(1)]=ecu.group(2)
# pprint.pprint(dict)

print('----------- 变速器型式/档位数：----------------')
shiftma_regex=re.compile(r'\d\d(变速器型式/档位数：)(.*)(\d\d)(消声器型号)')
shiftma=shiftma_regex.search(text)
# print(shiftma.group())
dict[shiftma.group(1)]=shiftma.group(2)
# pprint.pprint(dict)

print('----------- 消声器型号/生产企业：----------------')
voice_regex=re.compile(r'\d\d(消声器型号/生产企业：)(.*)(\d\d)(增压器型号)')
voice=voice_regex.search(text)
# print(voice.group())
dict[voice.group(1)]=voice.group(2)
pprint.pprint(dict)

print('----------- 增压器型号/生产企业：----------------')
charge_regex=re.compile(r'\d\d(增压器型号/生产企业：)(.*)(\d\d)(中冷器型式)')
charge=charge_regex.search(text)
# print(charge.group())
dict[charge.group(1)]=charge.group(2)
# pprint.pprint(dict)

print('----------- 中冷器型式：----------------')
zhonglen_regex=re.compile(r'\d\d(中冷器型式：)(.*)(标定标识1)')
zhonglen=zhonglen_regex.search(text)
# print(zhonglen.group())
dict[zhonglen.group(1)]=zhonglen.group(2)
# pprint.pprint(dict)
print(os.getcwd())

print('----------- cal1：----------------')
cal1_regex=re.compile(r'(标定标识1)(.*)(标定验证号1)')
cal1=cal1_regex.search(text)
# print(cal1.group())
dict[cal1.group(1)]=cal1.group(2)

print('----------- cvn1：----------------')
cvn1_regex=re.compile(r'(标定验证号1)(.*)(标定标识2)')
cvn1=cvn1_regex.search(text)
# print(cvn1.group())
dict[cvn1.group(1)]=cvn1.group(2)

print('----------- cal2：----------------')
cal2_regex=re.compile(r'(标定标识2)(.*)(标定验证号2)')
cal2=cal2_regex.search(text)
# print(cal2.group())
dict[cal2.group(1)]=cal2.group(2)

print('----------- cvn2：----------------')
cvn2_regex=re.compile(r'(标定验证号2)(.*)')
cvn2=cvn2_regex.search(text)
# print(cvn2.group())
dict[cvn2.group(1)]=cvn2.group(2)










cheliangtype=input('请输入车辆类型：')
zuidazhongliang=input('请输入最大设计重量：')
pailiang=input('请输入车辆排量：')
gangshu=input('请输入缸数：')

dict['车辆类型']=cheliangtype
dict['最大设计重量']=zuidazhongliang
dict['车辆排量']=pailiang
dict['缸数']=gangshu

pprint.pprint(dict)









def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
def get_paragraph(paras, text):
    for para in paras:
        if text in para.text:
            return para
    raise KeyError("The text cannot be found anywhere in the document")
if cheliangtype=='轿车':
    document=Document(os.getcwd()+'/jiaochemoban.docx')
if cheliangtype=='多用途乘用车':
    document=Document(os.getcwd()+'/chengyongchemoban.docx')
para1=get_paragraph(document.paragraphs,'123')
# table=document.add_table(row=12,cols=4)
# document.styles['Normal'].font.name = u'宋体'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


table = document.add_table(rows=12, cols=4)
a1=table.cell(0,0)
a2=table.cell(1,0)
A1=a1.merge(a2)
b1=table.cell(0,1)
b2=table.cell(1,1)
B1=b1.merge(b2)



b9=table.cell(8,1)
d9=table.cell(8,3)
B9=b9.merge(d9)


b10=table.cell(9,1)
d10=table.cell(9,3)
B10=b10.merge(d10)

b11=table.cell(10,1)
d11=table.cell(10,3)
B11=b11.merge(d11)

b12=table.cell(11,1)
d12=table.cell(11,3)
B12=b12.merge(d12)

# b13=table.cell(12,1)
# d13=table.cell(12,3)
# B13=b13.merge(d13)

A1.text='样品名称'
# A1.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
B1.text=cheliangtype
table.cell(0,2).text='型号规格'
table.cell(1,2).text='商    标'
table.cell(0,3).text=dict['车辆型号：']
table.cell(1,3).text=dict['商标']
# run = table.cell(0, 2).paragraphs[0].runs[0]
# run.font.bold = True
# run.font.size = Pt(14)

# for i in range(2,13):
#     for j in range(0,4):
#         table.cell(i,j).text=str(i)+','+str(j)
table.cell(2,0).text='委托单位'
table.cell(2,1).text='洛阳市机动车排气污染监控中心'
table.cell(2,2).text='检验类别'
table.cell(2,3).text='委托检验'

table.cell(3,0).text='生产单位'
table.cell(3,1).text=dict['车辆制造商名称：']
table.cell(3,2).text='样品等级'
table.cell(3,3).text='--'

table.cell(4,0).text='送样地点'
table.cell(4,1).text='--'
table.cell(4,2).text='样品等级'
table.cell(4,3).text='--'

table.cell(5,0).text='样品数量'
table.cell(5,1).text='一辆'
table.cell(5,2).text='送样者'
table.cell(5,3).text='--'

table.cell(6,0).text='抽样单位'
table.cell(6,1).text='国家机动车质量监督检验中心（重庆）'
table.cell(6,2).text='抽样者'
table.cell(6,3).text='李成果等'

table.cell(7,0).text='抽样基数'
table.cell(7,1).text='--'
table.cell(7,2).text='原编号或生产日期'
table.cell(7,3).text='--'

table.cell(8,0).text='检验依据'
table.cell(8,1).text='《汽油车污染物排放限值及测量方法（双怠速法及简易工况法）》'

table.cell(9,0).text='检验项目'
table.cell(9,1).text='外观检验、车载诊断系统（OBD）检查'

table.cell(10,0).text='检验结论'
table.cell(10,1).text='经检验，'+str(dict['车辆型号：'])+cheliangtype+'车型样品所检外观检验、车载诊断系统（OBD）检查项目的检验结果符合GB18285-2018《汽油车污染物排放限值及测量方法（双怠速法及简易工况法）》标准中的要求。\n          签发日期：    年  月  日        '

table.cell(11,0).text='备注'




move_table_after(table, para1)
# move_table_after(table, '123')
for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(12)
table.style = 'Table Grid'
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


#下面的部分是word报告中最后的一个表
#上面的部分是word报告中的第一个表格
table2=document.add_table(18,9)
table2.style='Table Grid'
table2.cell(0,0).merge(table2.cell(0,8))
table2.cell(0,0).text='3.1 基本信息'

table2.cell(1,0).text='车辆型号'
table2.cell(1,1).merge(table2.cell(1,2))
table2.cell(1,1).text=dict['车辆型号：']
table2.cell(1,3).merge(table2.cell(1,5))
table2.cell(1,3).text="车辆识别代号（VIN）"
table2.cell(1,6).merge(table2.cell(1,8))
table2.cell(1,6).text=dict['vin']

table2.cell(2,0).text='车辆生产企业'
table2.cell(2,1).merge(table2.cell(2,2))
table2.cell(2,1).text=dict['车辆制造商名称：']
table2.cell(2,3).merge(table2.cell(2,5))
table2.cell(2,3).text="车辆排放阶段"
table2.cell(2,6).merge(table2.cell(2,8))
table2.cell(2,6).text='国六'

table2.cell(3,0).text='变速箱型式'
table2.cell(3,1).merge(table2.cell(3,2))
table2.cell(3,1).text=dict['变速器型式/档位数：']
table2.cell(3,3).merge(table2.cell(3,5))
table2.cell(3,3).text="催化转化器型号"
table2.cell(3,6).merge(table2.cell(3,8))
table2.cell(3,6).text=dict['催化转化器型号']

table2.cell(4,0).text='基准质量（kg）'
table2.cell(4,1).merge(table2.cell(4,2))
table2.cell(4,1).text=dict['基准质量：']
table2.cell(4,3).merge(table2.cell(4,5))
table2.cell(4,3).text="最大总设计质量（kg）"
table2.cell(4,6).merge(table2.cell(4,8))
table2.cell(4,6).text=dict['最大设计重量']


table2.cell(5,0).text='发动机型号'
table2.cell(5,1).merge(table2.cell(5,2))
table2.cell(5,1).text=dict['发动机型号/生产企业：1']
table2.cell(5,3).merge(table2.cell(5,5))
table2.cell(5,3).text="发动机编号"
table2.cell(5,6).merge(table2.cell(5,8))
table2.cell(5,6).text=dict['发动机编号：']

table2.cell(6,0).text='发动机生产企业'
table2.cell(6,1).merge(table2.cell(6,2))
table2.cell(6,1).text=dict['发动机型号/生产企业：2']
table2.cell(6,3).merge(table2.cell(6,5))
table2.cell(6,3).text="发动机排量（L）"
table2.cell(6,6).merge(table2.cell(6,8))
table2.cell(6,6).text=pailiang

table2.cell(7,0).text='气缸数'
table2.cell(7,1).merge(table2.cell(7,2))
table2.cell(7,1).text=gangshu
table2.cell(7,3).merge(table2.cell(7,5))
table2.cell(7,3).text="燃油供给方式"
table2.cell(7,6).merge(table2.cell(7,8))
table2.cell(7,6).text='--'

table2.cell(8,0).text='电动机型号'
table2.cell(8,1).merge(table2.cell(8,2))
table2.cell(8,1).text='--'
table2.cell(8,3).merge(table2.cell(8,5))
table2.cell(8,3).text="储能装置型号"
table2.cell(8,6).merge(table2.cell(8,8))
table2.cell(8,6).text='--'

table2.cell(9,0).text='电池容量'
table2.cell(9,1).merge(table2.cell(9,2))
table2.cell(9,1).text='--'
table2.cell(9,3).merge(table2.cell(9,5))
table2.cell(9,3).text="OBD接口位置"
table2.cell(9,6).merge(table2.cell(9,8))
table2.cell(9,6).text='--'

table2.cell(10,0).merge(table2.cell(10,8))
table2.cell(10,0).text='3.2 外观检验'

table2.cell(11,0).merge(table2.cell(11,8))
table2.cell(11,0).text='本车实车污染控制装置与环保随车清单信息一致。'

table2.cell(12,0).merge(table2.cell(12,8))
table2.cell(12,0).text='3.3 OBD检查'

table2.cell(13,0).merge(table2.cell(13,3))
table2.cell(13,0).text='OBD通讯是否正常'

table2.cell(13,4).merge(table2.cell(13,8))
table2.cell(13,4).text='þ是  ¨ 否'

table2.cell(14,0).merge(table2.cell(16,1))
table2.cell(14,0).text='CAL ID/CVN信息'

table2.cell(14,2).merge(table2.cell(14,3))
table2.cell(14,2).text='发动机控制单元'

table2.cell(15,2).merge(table2.cell(15,3))
table2.cell(15,2).text='后处理控制单元（如适用）'

table2.cell(16,2).merge(table2.cell(16,3))
table2.cell(16,2).text='其他控制单元（如适用）'

table2.cell(14,4).text='CAL ID'
table2.cell(15,4).text='CAL ID'
table2.cell(16,4).text='CAL ID'

table2.cell(14,5).merge(table2.cell(14,6))
table2.cell(14,5).text=dict['标定标识1']

table2.cell(15,5).merge(table2.cell(15,6))
table2.cell(15,5).text='--'

table2.cell(16,5).merge(table2.cell(16,6))
table2.cell(16,5).text=dict['标定标识2']

table2.cell(14,7).text='CVN'
table2.cell(14,8).text=dict['标定验证号1']
table2.cell(15,7).text='CVN'
table2.cell(15,8).text="--"
table2.cell(16,7).text='CVN'
table2.cell(16,8).text=dict['标定验证号2']

table2.cell(17,0).merge(table2.cell(17,3))
table2.cell(17,0).text="OBD检查结果"

table2.cell(17,4).merge(table2.cell(17,8))
table2.cell(17,4).text='þ合格    ¨不合格'




for row in table2.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(12)
for row in table2.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
for row in table2.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
table2.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
table2.cell(10,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
table2.cell(12,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

para2=get_paragraph(document.paragraphs,'333')
move_table_after(table2, para2)
table3=document.add_table(12,4)
table3.style='Table Grid'

table3.cell(0,0).text='车辆型号'
table3.cell(0,1).text=dict['车辆型号：']
table3.cell(0,2).text='商 标 '
table3.cell(0,3).text=dict['商标']

table3.cell(1,0).text='汽车分类'
table3.cell(1,1).text=dict['汽车分类：']
table3.cell(1,2).text='排放阶段 '
table3.cell(1,3).text='国六'

table3.cell(2,0).text='车辆的识别方法和位置'
table3.cell(2,1).text='右侧B柱'
table3.cell(2,2).text='车辆制造商名称'
table3.cell(2,3).text=dict['车辆制造商名称：']

table3.cell(3,0).text='生产厂地址'
table3.cell(3,1).text=dict['生产厂地址：']
table3.cell(3,2).text='发动机编号  '
table3.cell(3,3).text=dict['发动机编号：']

table3.cell(4,0).text='基准质量（kg）'
table3.cell(4,1).text=dict['基准质量：']
table3.cell(4,2).text='发动机型号/生产企业  '
table3.cell(4,3).text=dict['发动机型号/生产企业：']

table3.cell(5,0).text='催化转化器型号/生产企业'
table3.cell(5,1).text=dict['催化转化器']
table3.cell(5,2).text='涂层/载体/封装生产企业  '
table3.cell(5,3).text=dict['催化转化器涂层']

table3.cell(6,0).text='颗粒捕集器型号/生产企业'
table3.cell(6,1).text=dict['颗粒捕集器']
table3.cell(6,2).text='涂层/载体/封装生产企业  '
table3.cell(6,3).text=dict['颗粒捕集器涂层']

table3.cell(7,0).text='炭罐型号/生产企业'
table3.cell(7,1).text=dict['炭罐型号/生产企业：']
table3.cell(7,2).text='氧传感器型号/生产企业  '
table3.cell(7,3).text=dict['氧传感器型号/生产企业：']

table3.cell(8,0).text='曲轴箱排放控制装置型号/生产企业'
table3.cell(8,1).text=dict['曲轴箱排放控制装置型号/生产企业：']
table3.cell(8,2).text='EGR型号/生产企业  '
table3.cell(8,3).text=dict['EGR型号/生产企业：']


table3.cell(9,0).text='OBD系统供应商'
table3.cell(9,1).text=dict['OBD系统供应商：']
table3.cell(9,2).text='ECU型号/生产企业  '
table3.cell(9,3).text=dict['ECU型号/生产企业：']

table3.cell(10,0).text='变速器型式/档位数'
table3.cell(10,1).text=dict['变速器型式/档位数：']
table3.cell(10,2).text='消声器型号/生产企业  '
table3.cell(10,3).text=dict['消声器型号/生产企业：']

table3.cell(11,0).text='增压器型号/生产企业'
table3.cell(11,1).text=dict['增压器型号/生产企业：']
table3.cell(11,2).text='中冷器型式  '
table3.cell(11,3).text=dict[ '中冷器型式：']

for row in table3.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(12)
for row in table3.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
for row in table3.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for row in table3.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(12)
para3=get_paragraph(document.paragraphs,'222')
move_table_after(table3, para3)
document.save('/Users/wangwang/PycharmProjects/docx_learning/'+str(dict['车辆型号：'])+'报告'+'.docx')
print(os.getcwd())



wb=openpyxl.load_workbook(os.getcwd()+'/pathtest.xlsx')
sheet=wb.get_sheet_by_name('参数')
sheet['D4'].value=cheliangxinghao.group(2)
sheet['H4'].value=brand.group(2)
sheet['D5'].value=qichefenlei.group(2)
sheet['H6'].value=oem.group(2)
sheet['D7'].value=address.group(2)
sheet['H7'].value=serial.group(2)
sheet['D8'].value=mass.group(2)
sheet['H8'].value=type_eng.group(2)
sheet['D9'].value=catalyst.group(1)
sheet['H9'].value=catalyst_layer.group(2)
sheet['D10'].value=dpf.group(2)
sheet['H10'].value=dpf_layer.group(4)
sheet['D11'].value=carbon.group(2)
sheet['H11'].value=ox.group(2)
sheet['D12'].value=crank.group(2)
sheet['H12'].value=egr.group(2)
sheet['D13'].value=obd.group(2)
sheet['H13'].value=ecu.group(2)
sheet['D14'].value=shiftma.group(2)
sheet['H14'].value=voice.group(2)
sheet['D15'].value=charge.group(2)
sheet['H15'].value=zhonglen.group(2)
sheet=wb.get_sheet_by_name('原始记录')
sheet['C4'].value=cheliangtype
sheet['G4'].value=cheliangxinghao.group(2)
sheet['C5'].value=vin.group(1)
sheet['D9'].value=type_eng1.group(2)
sheet['D11'].value=type_eng1.group(3)
print('处理完成')
wb.save(str(cheliangxinghao.group(2))+'原始记录.xlsx')

chexinglist=str(cheliangxinghao.group(2))
oem_name=str(oem.group(2))
with open('车辆型号清单.txt','a') as file_handle:
    file_handle.write(chexinglist)
    file_handle.write('\n')
with open('车辆制造商清单.txt','a',encoding='utf-8') as file_handle:
    file_handle.write(oem_name)
    file_handle.write('\n')


